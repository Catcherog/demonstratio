#!/usr/bin/env python3
"""Build and verify the canonical two-page Chinese AI/Agent PM resume.

This script intentionally uses only the Codex bundled Python runtime.  It is
kept beside the deliverable so the DOCX, TXT, PDF and render QA can be rebuilt
from the same canonical text.
"""
from __future__ import annotations

import argparse
import os
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[2]
PORTFOLIO_URL = "https://jaelchen-portfolio-vercel-extracted.vercel.app"
OUT = ROOT / "public" / "resume"
STEM = "chen-jiawei-ai-agent-cn-two-page"
DOCX = OUT / f"{STEM}.docx"
PDF = OUT / f"{STEM}.pdf"
TXT = OUT / f"{STEM}.txt"
PHOTO = OUT / "resume-portrait.png"
QA = OUT / "qa-render"
BLACK = "111111"
GRAY = "666666"
LIGHT = "F2F2F2"


CANONICAL_TEXT = """陈嘉伟 | AI / Agent 产品经理
杭州 | 18874988048 | Jael_Chen@foxmail.com | 作品集: https://jaelchen-portfolio-vercel-extracted.vercel.app | GitHub: Catcherog

专业摘要
具备复杂项目交付与 AI 产品实践经验，擅长把真实业务流程拆解为可验证的产品链路，覆盖数据治理、Agent 编排、多模态工具与人机协作边界。能够推进需求定义、原型验证、开发协同和迭代复盘，并明确区分验证、试点与生产状态。

工作经历
泽怀摄影工作室 | 创始人兼 AI 产品负责人 | 2026.02 - 至今
面向摄影业务规划客户触点、数据中台、智能服务与图像工具的产品协同路径；负责业务问题拆解、产品定义、原型验证与迭代推进。
围绕飞书 AI 业务数据平台、Service Agent 与光砚 Lumen 三个主案例，完成数据治理、Agent 工作流与图像产品化的产品规划及最小可用原型开发协同。

深圳市联洲国际技术有限公司（TP-Link） | 商用项目经理 | 2024.07 - 2026.02
负责 282 个 SKU 全生命周期，峰值并行推进 80+ 项目，覆盖交换机、PoE 供电、光模块 / OLT、Omada Controller 与多国特制软件五条产品线。
主导海外客户 NFC 功能需求定义，识别 Logo 合规与用户认知风险，提出“蓝牙 + NFC”融合交互方案并推动采纳。
在供应商更换和春节供应链中断风险下，建立并行验证与例外决策机制，追回 2 周工期，推动五款产品提前 15 天完成量产。

代表项目
飞书 AI 业务数据平台 | 数据治理与自动化
将聊天记录、截图、表单与人工录入转化为可治理的业务数据；设计 Candidate V1 合同、SOP BR-01~06 治理规则与审计链路。
真实 Tesseract OCR、三层幂等与真实测试 Base E2E 已验证；17 张表 / 12 条自动化仅为测试 Base 历史验收基线，不是 V2 自动化已上线证据；正式业务 Base 与生产上线未开放。

Service Agent | RAG 客服与知识飞轮
使用 LangGraph 8 节点 11 边工作流编排检索、风险分流、生成、质量检查与人工接管，完成 Web/API 端到端 MVP。
589 tests 全量回归（2026-07-28，非准确率）；固定 90 样本保留为静态 runner 审计输入。
现有 runner 未验证 expected_route、实际澄清节点或生成输出，因此不发布路由准确率、回答正确率或禁止承诺结果。
Controlled Demo｜公网前端可访问，后端恢复中；未接通时安全转人工。

光砚 Lumen | AI 图像编辑工作台
将多模型图像生成与编辑能力组织为含任务状态、版本与失败恢复的工作台，覆盖图片生成、编辑与多格式导出。
工程协作覆盖 Provider 抽象，以及项目 / 任务 / 结果持久化与恢复。
_id 更新缺陷已修复；本地 client/server 回归与 build 通过。
Controlled Demo｜后端健康已通过，真实核心编辑待有效凭据实测；采用 BYO Key 使用边界。

教育背景
中南大学 | 材料物理 | 本科 | 2020.09 - 2024.06
CET-6；基础 SQL；大学生创新创业项目省级奖项。

能力
AI 产品：Agent / RAG / 知识库 / 评估与人工兜底；产品交付：需求分析 / PRD / 原型 / 跨团队推进；技术基础：Python / TypeScript / SQL / API；工具与方法：LangGraph、飞书多维表、CloudBase、React / Expo、Prompt 与上下文设计。
"""

REQUIRED = [
    "AI / Agent 产品经理", "杭州", "282 个 SKU", "80+ 项目", "五条产品线",
    "追回 2 周", "五款产品提前 15 天", "飞书 AI 业务数据平台", "Service Agent",
    "光砚 Lumen", "测试 Base 历史验收基线", "不是 V2 自动化已上线证据", "589 tests",
    "固定 90 样本保留为静态 runner 审计输入", "未验证 expected_route", "不发布路由准确率、回答正确率或禁止承诺结果",
    "Controlled Demo", "公网前端可访问", "未接通时安全转人工", "Provider 抽象", "项目 / 任务 / 结果持久化与恢复",
    "_id 更新缺陷已修复", "本地 client/server 回归与 build 通过",
    "后端健康已通过", "真实核心编辑待有效凭据实测", "BYO Key",
]
FORBIDDEN = [
    "出生年份", "浙江省杭州市", "2002 年", "90%+", "92% 准确率", "50%+ 自动应答",
    "成熟生产", "生产级", "14 PASS + 1 PARTIAL", "实际离线评测待执行", "目前为 Online Beta",
    "75/90", "83.33%", "高风险错误放行 0", "禁止承诺 0",
]
ORDER = ["飞书 AI 业务数据平台", "Service Agent", "光砚 Lumen"]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, **edges: dict) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, values in edges.items():
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        for key, value in values.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_cell_margin(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run(run, size=9.2, bold=False, color=BLACK, font="Microsoft YaHei") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text: str, url: str, size=8.8) -> None:
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), BLACK); rpr.append(color)
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single"); rpr.append(underline)
    rfonts = OxmlElement("w:rFonts"); rfonts.set(qn("w:ascii"), "Arial"); rfonts.set(qn("w:hAnsi"), "Arial"); rfonts.set(qn("w:eastAsia"), "Microsoft YaHei"); rpr.append(rfonts)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2))); rpr.append(sz)
    run.append(rpr)
    txt = OxmlElement("w:t"); txt.text = text; run.append(txt)
    link.append(run); paragraph._p.append(link)


def paragraph(doc_or_cell, text="", size=9.2, bold=False, color=BLACK, before=0, after=2, align=None, style=None):
    p = doc_or_cell.add_paragraph(style=style) if style else doc_or_cell.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run(r, size, bold, color)
    return p


def section_heading(doc, number: str, title: str, english: str = "") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.line_spacing = 1
    r = p.add_run(f"{number}  {title}")
    set_run(r, 11.2, True, BLACK)
    if english:
        r = p.add_run(f"   {english.upper()}")
        set_run(r, 6.8, False, GRAY, "Arial")
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "12"); bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), BLACK)
    pbdr.append(bottom); p_pr.append(pbdr)


def add_bullet(doc, lead: str, body: str) -> None:
    p = doc.add_paragraph(style="Resume Bullet")
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.line_spacing = 1.13
    r = p.add_run(lead)
    set_run(r, 8.85, True)
    r = p.add_run(body)
    set_run(r, 8.85)


def add_project(doc, name: str, subtitle: str, bullets: list[tuple[str, str]]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(name)
    set_run(r, 10.1, True)
    r = p.add_run(f"  |  {subtitle}")
    set_run(r, 8.7, False, GRAY)
    for lead, body in bullets:
        add_bullet(doc, lead, body)


def footer(section, page: str) -> None:
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"陈嘉伟  |  AI / Agent 产品经理  |  {page}")
    set_run(r, 7.2, False, GRAY, "Arial")


def build_docx(photo_source: Path) -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    if photo_source and photo_source.exists():
        if photo_source.resolve() != PHOTO.resolve():
            shutil.copy2(photo_source, PHOTO)
    if not PHOTO.exists():
        raise FileNotFoundError("未找到证件照；请传入 --photo-source 或先放入 public/resume/resume-portrait.png")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.10)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)
    footer(section, "01 / 02")

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.2)
    bullet = styles.add_style("Resume Bullet", 1)
    bullet.base_style = normal
    bullet.paragraph_format.left_indent = Cm(0.42)
    bullet.paragraph_format.first_line_indent = Cm(-0.32)
    bullet.paragraph_format.space_after = Pt(2)
    bullet.paragraph_format.line_spacing = 1.13
    bullet.font.name = "Microsoft YaHei"
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # Header block: content uses a table for a stable photo and contact geometry.
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(12.8)
    table.columns[1].width = Cm(4.5)
    left, right = table.rows[0].cells
    for cell in (left, right):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margin(cell, 20, 0, 20, 0)
    set_cell_border(left, bottom={"val": "single", "sz": 16, "color": BLACK, "space": 4})
    set_cell_border(right, bottom={"val": "single", "sz": 16, "color": BLACK, "space": 4})
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("陈嘉伟")
    set_run(r, 23, True, BLACK)
    p = paragraph(left, "AI / Agent 产品经理", 11, True, BLACK, after=2)
    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run("杭州  |  18874988048  |  Jael_Chen@foxmail.com  |  ")
    set_run(r, 8.8)
    add_hyperlink(p, "作品集", PORTFOLIO_URL, 8.8)
    r = p.add_run("  |  GitHub: ")
    set_run(r, 8.8)
    add_hyperlink(p, "Catcherog", "https://github.com/Catcherog", 8.8)
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = rp.add_run()
    run.add_picture(str(PHOTO), width=Cm(3.0), height=Cm(3.6))

    section_heading(doc, "01", "专业摘要", "Profile")
    summary = doc.add_table(rows=1, cols=1)
    cell = summary.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    set_cell_margin(cell, 110, 140, 110, 140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run("具备复杂项目交付与 AI 产品实践经验，擅长把真实业务流程拆解为可验证的产品链路，覆盖数据治理、Agent 编排、多模态工具与人机协作边界。能够推进需求定义、原型验证、开发协同和迭代复盘，并明确区分验证、试点与生产状态。")
    set_run(r, 9.0)

    section_heading(doc, "02", "工作经历", "Experience")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("泽怀摄影工作室")
    set_run(r, 10, True)
    r = p.add_run("  |  创始人兼 AI 产品负责人  |  2026.02 - 至今")
    set_run(r, 8.6, False, GRAY)
    add_bullet(doc, "产品规划：", "面向摄影业务规划客户触点、数据中台、智能服务与图像工具的产品协同路径；负责业务问题拆解、产品定义、原型验证与迭代推进。")
    add_bullet(doc, "主案例协同：", "围绕飞书 AI 业务数据平台、Service Agent 与光砚 Lumen，完成数据治理、Agent 工作流与图像产品化的产品规划及最小可用原型开发协同。")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("深圳市联洲国际技术有限公司（TP-Link）")
    set_run(r, 10, True)
    r = p.add_run("  |  商用项目经理  |  2024.07 - 2026.02")
    set_run(r, 8.6, False, GRAY)
    add_bullet(doc, "多产品线交付：", "负责 282 个 SKU 全生命周期，峰值并行推进 80+ 项目，覆盖交换机、PoE 供电、光模块 / OLT、Omada Controller 与多国特制软件五条产品线。")
    add_bullet(doc, "跨国需求定义：", "主导海外客户 NFC 功能需求定义，识别 Logo 合规与用户认知风险，提出“蓝牙 + NFC”融合交互方案并推动采纳。")
    add_bullet(doc, "风险决策：", "在供应商更换和春节供应链中断风险下，建立并行验证与例外决策机制，追回 2 周工期，推动五款产品提前 15 天完成量产。")

    # Deliberate two-page structure: page one closes with the Feishu project.
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("代表项目")
    set_run(r, 18, True)
    r = p.add_run("  PROJECT HIGHLIGHTS")
    set_run(r, 7.2, False, GRAY, "Arial")
    section_heading(doc, "03", "飞书 AI 业务数据平台", "Data Governance")
    add_project(doc, "飞书 AI 业务数据平台", "数据治理与自动化", [
        ("目标：", "将聊天记录、截图、表单与人工录入转化为可治理的业务数据。"),
        ("设计：", "设计 Candidate V1 合同、SOP BR-01~06 治理规则与审计链路。"),
        ("历史基线：", "17 张表 / 12 条自动化仅为测试 Base 历史验收基线，不是 V2 自动化已上线证据。"),
        ("验证边界：", "真实 Tesseract OCR、三层幂等与真实测试 Base E2E 已验证；正式业务 Base 与生产上线未开放。"),
    ])

    second = doc.add_section(WD_SECTION.NEW_PAGE)
    second.page_width = Cm(21)
    second.page_height = Cm(29.7)
    second.top_margin = Cm(1.25)
    second.bottom_margin = Cm(1.10)
    second.left_margin = Cm(1.35)
    second.right_margin = Cm(1.35)
    second.footer_distance = Cm(0.55)
    second.footer.is_linked_to_previous = False
    footer(second, "02 / 02")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("代表项目 · 续")
    set_run(r, 18, True)
    r = p.add_run("  PROJECT HIGHLIGHTS")
    set_run(r, 7.2, False, GRAY, "Arial")
    section_heading(doc, "04", "Service Agent", "RAG Service")
    add_project(doc, "Service Agent", "RAG 客服与知识飞轮", [
        ("目标：", "使用 LangGraph 8 节点 11 边工作流编排检索、风险分流、生成、质量检查与人工接管，完成 Web/API 端到端 MVP。"),
        ("验证：", "589 tests 全量回归（2026-07-28，非准确率）；固定 90 样本保留为静态 runner 审计输入。"),
        ("评测边界：", "现有 runner 未验证 expected_route、实际澄清节点或生成输出，因此不发布路由准确率、回答正确率或禁止承诺结果。"),
        ("公网状态：", "Controlled Demo｜公网前端可访问，后端恢复中；未接通时安全转人工。"),
    ])
    section_heading(doc, "05", "光砚 Lumen", "Multimodal Workspace")
    add_project(doc, "光砚 Lumen", "AI 图像编辑工作台", [
        ("产品：", "将多模型图像生成与编辑能力组织为含任务状态、版本与失败恢复的工作台，覆盖图片生成、编辑与多格式导出。"),
        ("工程协作：", "Provider 抽象，以及项目 / 任务 / 结果持久化与恢复。"),
        ("本地验证：", "_id 更新缺陷已修复；本地 client/server 回归与 build 通过。"),
        ("状态与边界：", "Controlled Demo｜后端健康已通过，真实核心编辑待有效凭据实测；采用 BYO Key 使用边界。"),
    ])
    section_heading(doc, "06", "教育背景与能力", "Education & Skills")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("中南大学")
    set_run(r, 9.4, True)
    r = p.add_run("  |  材料物理  |  本科  |  2020.09 - 2024.06")
    set_run(r, 8.7, False, GRAY)
    p = paragraph(doc, "CET-6；基础 SQL；大学生创新创业项目省级奖项。", 8.8, False, BLACK, after=3)
    skills = doc.add_table(rows=2, cols=2)
    skills.autofit = False
    labels = [("AI 产品", "Agent / RAG / 知识库 / 评估与人工兜底"), ("产品交付", "需求分析 / PRD / 原型 / 跨团队推进"), ("技术基础", "Python / TypeScript / SQL / API"), ("工具与方法", "LangGraph、飞书多维表、CloudBase、React / Expo、Prompt 与上下文设计")]
    for cell, (label, text) in zip([skills.cell(0,0), skills.cell(0,1), skills.cell(1,0), skills.cell(1,1)], labels):
        set_cell_shading(cell, LIGHT)
        set_cell_margin(cell, 70, 90, 70, 90)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(label); set_run(r, 8.3, True)
        p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.06
        r = p.add_run(text); set_run(r, 7.8)

    doc.core_properties.author = "陈嘉伟"
    doc.core_properties.title = "陈嘉伟｜AI / Agent 产品经理"
    doc.core_properties.subject = "中文两页简历"
    doc.save(DOCX)
    TXT.write_text(CANONICAL_TEXT, encoding="utf-8")


def extract_docx_text(path: Path) -> str:
    doc = Document(path)
    blocks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                blocks.extend(p.text for p in cell.paragraphs)
    return "\n".join(blocks)


def docx_uris(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as z:
        rels = ET.fromstring(z.read("word/_rels/document.xml.rels"))
    return [node.attrib.get("Target", "") for node in rels if node.attrib.get("TargetMode") == "External"]


def pdf_uris(path: Path) -> list[str]:
    found = []
    reader = PdfReader(str(path))
    for page in reader.pages:
        for annot_ref in page.get("/Annots", []) or []:
            annot = annot_ref.get_object()
            action = annot.get("/A")
            if action and action.get("/URI"):
                found.append(str(action.get("/URI")))
    return found


def facts_check(text: str, expect_pass: bool) -> list[str]:
    errors = []
    for item in REQUIRED:
        if item not in text:
            errors.append(f"缺少必备事实：{item}")
    for item in FORBIDDEN:
        if item in text:
            errors.append(f"出现禁止内容：{item}")
    positions = [text.find(x) for x in ORDER]
    if any(p < 0 for p in positions) or positions != sorted(positions):
        errors.append("三项目顺序必须是：飞书 → Service Agent → 光砚 Lumen")
    if expect_pass and errors:
        raise AssertionError("；".join(errors))
    if not expect_pass and not errors:
        raise AssertionError("旧资产意外通过：校验器没有发现结构或事实问题")
    return errors


def verify_final() -> None:
    for artifact in (DOCX, PDF, TXT):
        if not artifact.exists() or artifact.stat().st_size == 0:
            raise AssertionError(f"缺少或为空：{artifact}")
    facts_check(TXT.read_text(encoding="utf-8"), True)
    facts_check(extract_docx_text(DOCX), True)
    reader = PdfReader(str(PDF))
    if len(reader.pages) != 2:
        raise AssertionError(f"PDF 必须恰好两页，实际为 {len(reader.pages)} 页")
    pdf_text = "\n".join(page.extract_text() or "" for page in reader.pages)
    facts_check(pdf_text, True)
    if not any("陈嘉伟" in (page.extract_text() or "") for page in reader.pages):
        raise AssertionError("PDF 中文文本不可提取")
    if PORTFOLIO_URL not in docx_uris(DOCX):
        raise AssertionError("DOCX 缺少已验证作品集超链接")
    if PORTFOLIO_URL not in pdf_uris(PDF):
        raise AssertionError("PDF 缺少已验证作品集超链接")
    print(f"PASS pages=2 pdf_bytes={PDF.stat().st_size} docx_links={docx_uris(DOCX)} pdf_links={pdf_uris(PDF)}")


def ensure_soffice_compat() -> None:
    """Create a tiny executable launcher only when LibreOffice is unavailable.

    The bundled renderer calls an executable named `soffice` directly.  This
    Windows image has Word but not LibreOffice, so the launcher translates that
    one conversion call to Word's PDF exporter without changing the renderer.
    """
    shim_dir = Path(__file__).parent
    exe = shim_dir / "soffice.exe"
    source = shim_dir / "soffice_compat.cs"
    if exe.exists() and exe.stat().st_mtime >= source.stat().st_mtime:
        return
    compiler = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Microsoft.NET" / "Framework64" / "v4.0.30319" / "csc.exe"
    if not compiler.exists():
        raise FileNotFoundError("未找到 LibreOffice 或 C# 兼容启动器编译器")
    subprocess.run([str(compiler), "/nologo", f"/out:{exe}", str(source)], check=True)


def ensure_poppler_compat() -> None:
    shim_dir = Path(__file__).parent
    source = shim_dir / "poppler_compat.cs"
    compiler = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Microsoft.NET" / "Framework64" / "v4.0.30319" / "csc.exe"
    if not compiler.exists():
        raise FileNotFoundError("未找到 Poppler 兼容启动器编译器")
    for tool in ("pdfinfo", "pdftoppm"):
        exe = shim_dir / f"{tool}.exe"
        if not exe.exists() or exe.stat().st_mtime < source.stat().st_mtime:
            subprocess.run([str(compiler), "/nologo", f"/out:{exe}", str(source)], check=True)


def main() -> int:
    parser = argparse.ArgumentParser()
    sub = parser.add_subparsers(dest="cmd", required=True)
    old = sub.add_parser("check-old", help="Confirm legacy source fails the current fact/structure gate.")
    old.add_argument("path", type=Path)
    build = sub.add_parser("build", help="Build editable DOCX and canonical TXT.")
    build.add_argument("--photo-source", type=Path, default=PHOTO)
    all_cmd = sub.add_parser("all", help="Build, render to PDF/PNGs, and run all gates.")
    all_cmd.add_argument("--photo-source", type=Path, default=PHOTO)
    all_cmd.add_argument("--renderer", type=Path, required=True)
    args = parser.parse_args()
    try:
        if args.cmd == "check-old":
            errors = facts_check(args.path.read_text(encoding="utf-8"), False)
            print("EXPECTED_FAIL " + " | ".join(errors))
            return 1
        if args.cmd == "build":
            build_docx(args.photo_source)
            facts_check(TXT.read_text(encoding="utf-8"), True)
            facts_check(extract_docx_text(DOCX), True)
            print(f"BUILT {DOCX} {TXT}")
            return 0
        build_docx(args.photo_source)
        if QA.exists():
            shutil.rmtree(QA)
        # The documents renderer invokes `soffice`.  Prefer a real LibreOffice
        # installation; on this Windows workspace the adjacent compatibility
        # launcher delegates to the locally installed Word PDF engine.
        os.environ["PATH"] = str(Path(__file__).parent) + os.pathsep + os.environ.get("PATH", "")
        os.environ["CODEX_RESUME_PYTHON"] = sys.executable
        ensure_soffice_compat()
        ensure_poppler_compat()
        subprocess.run([sys.executable, str(args.renderer), str(DOCX), "--output_dir", str(QA), "--emit_pdf"], check=True)
        generated_pdf = QA / f"{STEM}.pdf"
        if not generated_pdf.exists():
            raise AssertionError(f"渲染器未生成 PDF：{generated_pdf}")
        shutil.copy2(generated_pdf, PDF)
        verify_final()
        return 0
    except Exception as exc:
        print(f"FAIL {exc}", file=sys.stderr)
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
